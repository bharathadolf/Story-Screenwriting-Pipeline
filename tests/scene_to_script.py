import json
import os
import argparse
import docx
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ANSI Color Codes
class Colors:
    HEADER = '\033[95m'      # Magenta
    SCENE_HEADING = '\033[94m' # Blue
    CHARACTER = '\033[92m'   # Green
    PARENTHETICAL = '\033[93m' # Yellow
    DIALOGUE = '\033[97m'    # White
    ACTION = '\033[96m'      # Cyan
    ENDC = '\033[0m'         # Reset color

def setup_document():
    doc = docx.Document()
    
    # 1. Page Setup (A4)
    # Width: 21.0 cm, Height: 29.7 cm
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    
    # 2. Margins
    # Top: 2.5 cm, Bottom: 2.5 cm, Left: 3.8 cm, Right: 2.5 cm
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3.8)
    section.right_margin = Cm(2.5)
    
    # 3. Base Font Settings
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Courier'
    font.size = Pt(12)
    # Single line spacing, no spacing after
    style.paragraph_format.space_after = Pt(0)
    style.paragraph_format.line_spacing = 1.0

    # 4. Element Formal Dimensions
    
    # Scene Heading - 0cm, ALL CAPS
    style_scene = doc.styles.add_style('Scene Heading', docx.enum.style.WD_STYLE_TYPE.PARAGRAPH)
    style_scene.base_style = doc.styles['Normal']
    style_scene.paragraph_format.left_indent = Cm(0)
    style_scene.font.all_caps = True
    style_scene.paragraph_format.space_before = Pt(12) # spacing before the scene heading

    # Action - 0cm
    style_action = doc.styles.add_style('Action', docx.enum.style.WD_STYLE_TYPE.PARAGRAPH)
    style_action.base_style = doc.styles['Normal']
    style_action.paragraph_format.left_indent = Cm(0)
    style_action.paragraph_format.space_before = Pt(12)

    # Character - 6.3cm left
    style_char = doc.styles.add_style('Character', docx.enum.style.WD_STYLE_TYPE.PARAGRAPH)
    style_char.base_style = doc.styles['Normal']
    style_char.paragraph_format.left_indent = Cm(6.3)
    style_char.font.all_caps = True
    style_char.paragraph_format.space_before = Pt(12)
    style_char.paragraph_format.keep_with_next = True

    # Dialogue - 4.4cm left, 3.8cm right
    style_dialogue = doc.styles.add_style('Dialogue', docx.enum.style.WD_STYLE_TYPE.PARAGRAPH)
    style_dialogue.base_style = doc.styles['Normal']
    style_dialogue.paragraph_format.left_indent = Cm(4.4)
    style_dialogue.paragraph_format.right_indent = Cm(3.8)

    # Parenthetical - 5.4cm left, 4.6cm right
    style_paren = doc.styles.add_style('Parenthetical', docx.enum.style.WD_STYLE_TYPE.PARAGRAPH)
    style_paren.base_style = doc.styles['Normal']
    style_paren.paragraph_format.left_indent = Cm(5.4)
    style_paren.paragraph_format.right_indent = Cm(4.6)
    style_paren.paragraph_format.keep_with_next = True

    # Transition - Right aligned
    style_trans = doc.styles.add_style('Transition', docx.enum.style.WD_STYLE_TYPE.PARAGRAPH)
    style_trans.base_style = doc.styles['Normal']
    style_trans.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    style_trans.font.all_caps = True
    style_trans.paragraph_format.space_before = Pt(12)

    return doc

def generate_script_from_scene(scene_file_path, output_dir="."):
    """
    Simulates the 'Drafting Department' agentic workflow.
    Takes a structured scene data packet and renders it into a docx script format.
    """
    print(f"Loading scene data from {scene_file_path}...")
    with open(scene_file_path, 'r') as f:
        scene_data = json.load(f)
        
    scene_number = scene_data.get("scene_number", 0)
    location = scene_data.get("location", "UNKNOWN LOCATION")
    time = scene_data.get("time", "DAY")
    
    # Render the fountain/script text format with ANSI colors for terminal
    terminal_content = f"{Colors.HEADER}SCENE START{Colors.ENDC}\n\n"
    terminal_content += f"{Colors.SCENE_HEADING}EXT/INT. {location} - {time}{Colors.ENDC}\n\n"

    doc = setup_document()
    doc.add_paragraph(f"EXT/INT. {location} - {time}", style='Scene Heading')
    
    for action in scene_data.get("action_lines", []):
        terminal_content += f"{Colors.ACTION}{action}{Colors.ENDC}\n\n"
        doc.add_paragraph(action, style='Action')
        
    for line in scene_data.get("dialogue", []):
        char = line.get("character")
        paren = line.get("parenthetical", "")
        text = line.get("text")
        
        terminal_content += f"{Colors.CHARACTER}          {char}{Colors.ENDC}\n"
        doc.add_paragraph(char, style='Character')
        
        if paren:
            if not paren.startswith("("): paren = f"({paren})"
            if not paren.endswith(")"): paren = f"{paren})"
            terminal_content += f"{Colors.PARENTHETICAL}        {paren}{Colors.ENDC}\n"
            doc.add_paragraph(paren, style='Parenthetical')
            
        terminal_content += f"{Colors.DIALOGUE}    {text}{Colors.ENDC}\n\n"
        doc.add_paragraph(text, style='Dialogue')
        
    output_filename = f"scene_{scene_number}.docx"
    output_path = os.path.join(output_dir, output_filename)
    
    doc.save(output_path)
        
    print(f"\n+++ Successfully generated script file: {output_path} +++\n")
    print("--- PREVIEW OF COMPILED SCENE ---")
    print(terminal_content)

if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Convert scene JSON to script DOCX.")
    parser.add_argument("--input", type=str, help="Path to scene JSON file", default="scene_demo.json")
    parser.add_argument("--output_dir", type=str, help="Directory to save the script", default=".")
    
    args = parser.parse_args()
    
    # Ensure relative paths are handled correctly if run from outside the directory
    current_dir = os.path.dirname(os.path.abspath(__file__))
    input_path = os.path.join(current_dir, args.input) if not os.path.isabs(args.input) else args.input
    output_dir = os.path.join(current_dir, args.output_dir) if not os.path.isabs(args.output_dir) else args.output_dir
    
    if os.path.exists(input_path):
        generate_script_from_scene(input_path, output_dir)
    else:
        print(f"Error: Demo scene file not found at {input_path}")
